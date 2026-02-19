"""
Agent 4: Gap Audit + Report Generator
=======================================
Compares available data against ESRS E1 (Climate) disclosure requirements.
Generates a professional Word document report with full data citations.

ALL reports are marked DRAFT — REQUIRES HUMAN REVIEW.
"""
import os
from datetime import datetime, date
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from rich.console import Console
from rich.table import Table as RichTable

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from db import (
    get_session, log_audit,
    Company, AccountMapping, Transaction, EmissionRecord,
    Supplier, EsrsDisclosure, ReportVersion,
)

console = Console()

# ═══════════════════════════════════════════════════════════════
# ESRS E1 DISCLOSURE REQUIREMENTS
# European Sustainability Reporting Standards — Climate
# ═══════════════════════════════════════════════════════════════
ESRS_E1_REQUIREMENTS = [
    {
        "ref": "E1-1",
        "title": "Transition plan for climate change mitigation",
        "data_needed": "Strategic climate targets, decarbonisation pathway, capex plans",
        "data_source": "management_input",
        "auto_assessable": False,
    },
    {
        "ref": "E1-2",
        "title": "Policies related to climate change mitigation and adaptation",
        "data_needed": "Documented climate/energy policies",
        "data_source": "management_input",
        "auto_assessable": False,
    },
    {
        "ref": "E1-3",
        "title": "Actions and resources related to climate change",
        "data_needed": "Climate actions taken, resources allocated",
        "data_source": "management_input",
        "auto_assessable": False,
    },
    {
        "ref": "E1-4",
        "title": "Targets related to climate change mitigation and adaptation",
        "data_needed": "GHG reduction targets with base year and timeline",
        "data_source": "management_input",
        "auto_assessable": False,
    },
    {
        "ref": "E1-5",
        "title": "Energy consumption and mix",
        "data_needed": "Total energy consumption in MWh, breakdown by source (electricity, gas, fuel, renewable share)",
        "data_source": "bmd_fibu",
        "auto_assessable": True,
        "check_categories": ["energy_electricity", "energy_gas"],
    },
    {
        "ref": "E1-6",
        "title": "Gross Scopes 1, 2, 3 and Total GHG emissions",
        "data_needed": "Scope 1/2/3 GHG emissions in tCO2e with calculation methodology",
        "data_source": "calculation_engine",
        "auto_assessable": True,
        "check_scopes": [1, 2, 3],
    },
    {
        "ref": "E1-7",
        "title": "GHG removals and GHG mitigation projects",
        "data_needed": "Carbon offsets purchased, removal credits, nature-based solutions",
        "data_source": "management_input",
        "auto_assessable": False,
    },
    {
        "ref": "E1-8",
        "title": "Internal carbon pricing",
        "data_needed": "Internal carbon price applied to investment decisions",
        "data_source": "management_input",
        "auto_assessable": False,
    },
    {
        "ref": "E1-9",
        "title": "Anticipated financial effects from material physical and transition risks",
        "data_needed": "Financial impact assessment of climate risks and opportunities",
        "data_source": "management_input",
        "auto_assessable": False,
    },
]


def run_gap_audit(company_id: int) -> list:
    """
    Assess each ESRS E1 requirement against available data.
    Returns list of disclosure assessments.
    """
    session = get_session()

    # Check what data we have
    has_mappings = session.query(AccountMapping).filter(
        AccountMapping.company_id == company_id
    ).count() > 0

    emissions = session.query(EmissionRecord).filter(
        EmissionRecord.company_id == company_id
    ).all()
    emission_scopes = {e.scope for e in emissions}
    emission_categories = {e.category for e in emissions}

    mapped_categories = set()
    if has_mappings:
        mappings = session.query(AccountMapping).filter(
            AccountMapping.company_id == company_id
        ).all()
        mapped_categories = {m.esg_category for m in mappings}

    supplier_count = session.query(Supplier).filter(
        Supplier.company_id == company_id
    ).count()

    assessments = []

    for req in ESRS_E1_REQUIREMENTS:
        status = "gap"
        notes = ""

        if req["auto_assessable"]:
            if "check_categories" in req:
                found = mapped_categories & set(req["check_categories"])
                if found and emissions:
                    status = "met"
                    notes = f"Data available from BMD FIBU. Categories mapped: {', '.join(found)}."
                elif found:
                    status = "partial"
                    notes = f"Account mappings exist ({', '.join(found)}) but emissions not yet calculated."
                else:
                    status = "gap"
                    notes = f"Required categories ({', '.join(req['check_categories'])}) not found in account mappings."

            elif "check_scopes" in req:
                found_scopes = emission_scopes & set(req["check_scopes"])
                if len(found_scopes) >= 2:
                    status = "met" if 3 in found_scopes else "partial"
                    notes = f"Emissions calculated for Scope(s) {', '.join(str(s) for s in sorted(found_scopes))}."
                    if 3 not in found_scopes:
                        notes += " Scope 3 data incomplete."
                elif found_scopes:
                    status = "partial"
                    notes = f"Only Scope {', '.join(str(s) for s in found_scopes)} available."
                else:
                    status = "gap"
                    notes = "No emissions data calculated yet."
        else:
            status = "gap"
            notes = f"Requires management input. Data source: {req['data_source']}. Cannot be auto-assessed from accounting data."

        # Save to database
        existing = session.query(EsrsDisclosure).filter(
            EsrsDisclosure.company_id == company_id,
            EsrsDisclosure.standard_ref == req["ref"],
        ).first()

        if existing:
            existing.status = status
            existing.gap_notes = notes
            existing.data_available = (status in ("met", "partial"))
            existing.last_assessed_at = datetime.utcnow()
        else:
            disc = EsrsDisclosure(
                company_id=company_id,
                standard_ref=req["ref"],
                disclosure_title=req["title"],
                status=status,
                data_available=(status in ("met", "partial")),
                gap_notes=notes,
                last_assessed_at=datetime.now(),
            )
            session.add(disc)

        assessments.append({
            "ref": req["ref"],
            "title": req["title"],
            "status": status,
            "notes": notes,
            "data_source": req["data_source"],
        })

    session.commit()
    session.close()
    return assessments


def generate_report(company_id: int, output_dir: str = None) -> str:
    """
    Generate a Word document ESRS E1 gap report with emissions data.
    Returns the file path of the generated report.
    """
    console.print(f"\n[bold blue]Agent 4: Gap Audit + Report Generation[/bold blue]")

    session = get_session()
    company = session.get(Company, company_id)
    if not company:
        console.print("[red]Company not found[/red]")
        session.close()
        return None

    # Store company info before session operations
    company_name = company.name
    company_uid = company.uid_vat

    # Run gap audit
    assessments = run_gap_audit(company_id)

    # Re-open session (gap audit may have closed it)
    session = get_session()

    # Get emissions data
    emissions = session.query(EmissionRecord).filter(
        EmissionRecord.company_id == company_id
    ).all()

    scope_totals = defaultdict(float)
    emission_details = []
    for e in emissions:
        scope_totals[e.scope] += e.value_tco2e
        emission_details.append({
            "scope": e.scope,
            "category": e.category,
            "value_tco2e": e.value_tco2e,
            "quantity": e.quantity,
            "unit": e.unit,
            "emission_factor_used": e.emission_factor_used,
            "factor_source": e.factor_source,
        })

    # Get mapping stats
    mapping_count = session.query(AccountMapping).filter(
        AccountMapping.company_id == company_id
    ).count()
    tx_count = session.query(Transaction).filter(
        Transaction.company_id == company_id
    ).count()

    # Determine period
    period = emissions[0].period if emissions else "N/A"
    session.close()

    # ─── Build Word Document ───
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ─── Cover Page ───
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ESRS E1 Climate Disclosure")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Gap Analysis & Emissions Report")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 92, 139)

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Company: {company_name}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(51, 51, 51)

    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info2.add_run(f"Period: {period}  |  Generated: {date.today().isoformat()}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")

    draft = doc.add_paragraph()
    draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = draft.add_run("DRAFT \u2014 REQUIRES HUMAN REVIEW")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(198, 40, 40)
    run.bold = True

    doc.add_page_break()

    # ─── Executive Summary ───
    doc.add_heading("1. Executive Summary", level=1)

    met = sum(1 for a in assessments if a["status"] == "met")
    partial = sum(1 for a in assessments if a["status"] == "partial")
    gap = sum(1 for a in assessments if a["status"] == "gap")
    total = scope_totals[1] + scope_totals[2] + scope_totals[3]

    doc.add_paragraph(
        f"This report assesses {company_name}'s data availability against the "
        f"ESRS E1 (Climate) disclosure requirements. The assessment is based on "
        f"{tx_count:,} financial transactions parsed from BMD NTCS, with "
        f"{mapping_count} accounts mapped to ESG categories."
    )

    p = doc.add_paragraph()
    p.add_run("Compliance Status: ").bold = True
    p.add_run(f"{met} of {len(assessments)} disclosures met, "
              f"{partial} partially met, {gap} gaps remaining.")

    if total > 0:
        p2 = doc.add_paragraph()
        p2.add_run("Total Estimated Emissions: ").bold = True
        p2.add_run(f"{total:,.2f} tCO2e "
                    f"(Scope 1: {scope_totals[1]:,.2f}, "
                    f"Scope 2: {scope_totals[2]:,.2f}, "
                    f"Scope 3: {scope_totals[3]:,.2f})")

    doc.add_paragraph(
        "Note: Quantities have been estimated from EUR spend using average Austrian "
        "energy prices. For auditable ESRS reports, these estimates should be replaced "
        "with actual consumption data from utility bills and fuel receipts."
    ).italic = True

    # ─── Gap Analysis ───
    doc.add_page_break()
    doc.add_heading("2. ESRS E1 Gap Analysis", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = table.rows[0].cells
    headers[0].text = "Reference"
    headers[1].text = "Disclosure Requirement"
    headers[2].text = "Status"
    headers[3].text = "Notes"

    for cell in headers:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for a in assessments:
        row = table.add_row().cells
        row[0].text = a["ref"]
        row[1].text = a["title"]

        status_text = a["status"].upper()
        row[2].text = status_text
        for paragraph in row[2].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                if a["status"] == "met":
                    run.font.color.rgb = RGBColor(46, 125, 50)
                elif a["status"] == "partial":
                    run.font.color.rgb = RGBColor(230, 81, 0)
                else:
                    run.font.color.rgb = RGBColor(198, 40, 40)

        row[3].text = a["notes"]

        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # ─── Emissions Detail ───
    if emission_details:
        doc.add_page_break()
        doc.add_heading("3. Emissions Calculation Detail", level=1)

        doc.add_paragraph(
            "All emissions are calculated using deterministic formulas. No AI/LLM was used "
            "for any numerical calculation. Each figure below is fully traceable to its "
            "source data, emission factor, and calculation methodology."
        )

        for scope_num in [1, 2, 3]:
            scope_emissions = [e for e in emission_details if e["scope"] == scope_num]
            if not scope_emissions:
                continue

            doc.add_heading(f"3.{scope_num} Scope {scope_num} Emissions", level=2)

            etable = doc.add_table(rows=1, cols=5)
            etable.style = "Light Grid Accent 1"

            eh = etable.rows[0].cells
            eh[0].text = "Category"
            eh[1].text = "Quantity"
            eh[2].text = "Factor"
            eh[3].text = "tCO2e"
            eh[4].text = "Source"

            for cell in eh:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)

            for e in scope_emissions:
                row = etable.add_row().cells
                row[0].text = e["category"] or ""
                row[1].text = f"{e['quantity']:,.1f} {e['unit']}" if e["quantity"] else "N/A"
                row[2].text = f"{e['emission_factor_used']} kg CO2e/{e['unit']}" if e["emission_factor_used"] else ""
                row[3].text = f"{e['value_tco2e']:,.4f}"
                row[4].text = e["factor_source"] or ""

                for cell in row:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

            subtotal = sum(e["value_tco2e"] for e in scope_emissions)
            trow = etable.add_row().cells
            trow[0].text = f"Scope {scope_num} Total"
            trow[3].text = f"{subtotal:,.4f}"
            for cell in [trow[0], trow[3]]:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)

    # ─── Methodology ───
    doc.add_page_break()
    doc.add_heading("4. Methodology & Data Sources", level=1)

    doc.add_heading("4.1 Data Collection", level=2)
    doc.add_paragraph(
        f"Financial data was extracted from BMD NTCS via CSV export "
        f"({tx_count:,} transactions processed). Each transaction was parsed with "
        f"encoding detection (UTF-8/Windows-1252), validated against BMD NTCS schema, "
        f"and stored with full source traceability (file name, row number, ingestion timestamp)."
    )

    doc.add_heading("4.2 Account Classification", level=2)
    doc.add_paragraph(
        f"{mapping_count} accounts were classified into ESG categories using "
        f"AI-assisted classification (Anthropic Claude API) with human confirmation. "
        f"Each account mapping was reviewed and confirmed by a qualified operator."
    )

    doc.add_heading("4.3 Emission Factors", level=2)
    doc.add_paragraph(
        "All emission factors are sourced from Umweltbundesamt Austria (vintage 2024). "
        "Calculations use the formula: quantity (physical units) x emission factor = kg CO2e. "
        "Where physical quantities were not directly available, they were estimated from "
        "EUR spend using average Austrian energy prices. These spend-based estimates should "
        "be validated against actual consumption data for auditable reports."
    )

    doc.add_heading("4.4 Limitations", level=2)
    doc.add_paragraph(
        "This report covers ESRS E1 (Climate) only. Disclosures E2-E5 (pollution, water, "
        "biodiversity, circular economy), S1-S4 (social), and G1 (governance) are not "
        "assessed in this version. Scope 3 emissions are limited to categories with "
        "available data (logistics/freight). Full Scope 3 assessment requires supplier "
        "outreach and additional data collection."
    )

    # ─── Audit Trail ───
    doc.add_heading("4.5 Audit Trail", level=2)
    doc.add_paragraph(
        "Every figure in this report is traceable through the following chain: "
        "Report figure \u2192 emissions_records table (calculation method, factor source) "
        "\u2192 transactions table (source_file, row_number) \u2192 original BMD CSV file. "
        "Full audit logs are maintained in the platform database."
    )

    # ─── Footer on every page ───
    doc.add_page_break()
    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(
        "This report has been generated automatically and is marked as DRAFT. "
        "It requires review and approval by a qualified person before submission. "
        "The platform does not provide legal advice or assurance of CSRD compliance. "
        "All emission calculations are estimates based on financial data and should be "
        "validated against actual consumption records for auditable reporting."
    )
    doc.add_paragraph(
        f"Generated: {datetime.now().isoformat()} | Company: {company_name} | "
        f"UID: {company_uid or 'N/A'}"
    ).italic = True

    # ─── Save ───
    if not output_dir:
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output")
    os.makedirs(output_dir, exist_ok=True)

    safe_period = period.replace("/", "-").replace("\\", "-")
    filename = f"ESRS_E1_Report_{company_name.replace(' ', '_')}_{safe_period}_{date.today().isoformat()}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    # Record report version
    session2 = get_session()
    existing_count = session2.query(ReportVersion).filter(
        ReportVersion.company_id == company_id,
        ReportVersion.report_type == "esrs_e1_gap",
    ).count()

    report = ReportVersion(
        company_id=company_id,
        report_type="esrs_e1_gap",
        version_number=existing_count + 1,
        status="draft",
        file_path=filepath,
    )
    session2.add(report)
    log_audit(session2, "report_generated", "report", filepath,
              f"type=esrs_e1_gap, version={existing_count + 1}, "
              f"met={met}, partial={partial}, gap={gap}")
    session2.commit()
    session2.close()

    # Print summary
    console.print(f"\n  [bold green]Report generated: {filepath}[/bold green]")

    summary_table = RichTable(title="ESRS E1 Gap Summary")
    summary_table.add_column("Status", style="bold")
    summary_table.add_column("Count", justify="center")
    summary_table.add_row("[green]MET[/green]", str(met))
    summary_table.add_row("[yellow]PARTIAL[/yellow]", str(partial))
    summary_table.add_row("[red]GAP[/red]", str(gap))
    console.print(summary_table)

    session.close()
    return filepath


if __name__ == "__main__":
    from db import init_db
    init_db()

    session = get_session()
    company = session.query(Company).first()
    session.close()

    if company:
        generate_report(company.id)
    else:
        console.print("[red]No company found. Run agents 1-3 first.[/red]")
